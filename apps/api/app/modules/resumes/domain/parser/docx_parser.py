import io
import xml.etree.ElementTree as ET
import zipfile

from apps.api.app.core.exceptions import DomainException
from apps.api.app.core.logging import logger
from apps.api.app.modules.resumes.domain.parser.provider import DocumentParserProvider


class DOCXParser(DocumentParserProvider):
    """
    Production DOCX text extraction parser.
    Uses docx library or secure zipfile XML element parsing for word/document.xml.
    Prevents macro execution and external XML entity resolution (XXE safe).
    """

    def parse_document(self, file_bytes: bytes) -> str:
        if not file_bytes:
            raise DomainException("DOCX file content is empty", code="EMPTY_DOCUMENT")

        try:
            # Try python-docx if installed
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            full_text = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                    if row_text:
                        full_text.append(row_text)
            return "\n".join(full_text).strip()
        except DomainException:
            raise
        except Exception:
            # Secure zipfile XML fallback parsing word/document.xml
            return self._parse_zip_xml(file_bytes)

    def _parse_zip_xml(self, file_bytes: bytes) -> str:
        try:
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
                if "word/document.xml" not in z.namelist():
                    raise DomainException("Invalid DOCX format: word/document.xml missing", code="MALFORMED_DOCX")
                xml_content = z.read("word/document.xml")

            root = ET.fromstring(xml_content)
            # Find all text elements <w:t>
            texts = []
            for elem in root.iter():
                if elem.tag.endswith("}t") or elem.tag == "t":
                    if elem.text:
                        texts.append(elem.text)
            return "\n".join(texts).strip()
        except DomainException:
            raise
        except Exception as e:
            logger.error(f"[DOCXParser] Zip XML parsing failed: {str(e)}")
            raise DomainException(f"Failed to parse DOCX file structure: {str(e)}", code="MALFORMED_DOCX")
